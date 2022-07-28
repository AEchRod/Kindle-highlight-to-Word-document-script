import requests
from docx import Document
# import datetime

document = Document()

# getting books that were updated last week, if needed:
# a_week_ago = datetime.datetime.now() - datetime.timedelta(days=7)

querystring_category = {
    "category": "books",
}

# To check if GET request works
response_auth = requests.get(
    url="https://readwise.io/api/v2/auth/",
    headers={"Authorization": "Token xxx"}
)

response_books = requests.get(
    url="https://readwise.io/api/v2/books/",
    headers={"Authorization": "Token xxx"},
    params=querystring_category
)

# this variable contains the whole booklist as a dictionary
data_booklist = response_books.json()

# this variable contains only the results from the booklist, since data_booklist['results'] is a list of dictionaries.
booklist_results = data_booklist['results']

# this variable creates a list of each of the book ids
book_x = [book['id'] for book in booklist_results]

book_title = [book['title'] for book in booklist_results]

# this creates a dictionary with both, the id and book title.
id_title = dict(zip(book_x, book_title))

# we unpack dictionary so we can get book highlights by book id and add titles to our file to be exported.
for id, title in id_title.items():

    querystring_bookid = {
        "book_id": id,
    }
    # this query string selects books individually

    # this requests gets the highlist list
    response = requests.get(
            url="https://readwise.io/api/v2/highlights/",
            headers={"Authorization": "Token xxx"},
            params=querystring_bookid
        )

    # this retrieves a json from the highlist list
    data_highlights = response.json()

    # this variable contains only the results from the booklist, data_highlight['results'] is a list of dictionaries
    highlight_results = data_highlights['results']

    # this makes the book titles bold by using the add.run() method.
    para = document.add_paragraph()
    bold_para = para.add_run(title).bold=True

    for highlight in highlight_results:
        document.add_paragraph(highlight['text'])
    document.save('Highlights.docx')
